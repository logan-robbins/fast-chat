"""
File processing service with smart routing for different document types.

This module implements intelligent file type routing:
- Text files (.txt, .md, .csv): Direct text extraction (UTF-8 decode)
- Structured data (.xlsx, .xls): Pandas extraction (no vision needed)
- Word documents (.docx): python-docx extraction (no vision needed)
- PDFs: Smart routing - try pdfplumber text extraction first, fall back to vision
- Presentations (.pptx): Vision extraction (layout matters)

Processing Strategy Selection:
1. TEXT_ONLY: Direct read or library extraction (fastest, no API calls)
2. STRUCTURED_DATA: Pandas DataFrame conversion (tabular data)
3. VISION_REQUIRED: GPT-4o vision API (scanned docs, presentations)

Threading:
    All synchronous extraction libraries (pdfplumber, python-docx, pandas)
    are wrapped in asyncio.to_thread() so they run in the default thread pool
    executor and do not block the asyncio event loop. This is especially
    important for large PDFs and Excel files which can be CPU-intensive.

SDK Versions (verified 02/05/2026):
- pdfplumber 0.11.9: PDF text extraction (synchronous, thread-offloaded)
- python-docx>=1.1.0: DOCX paragraph/table extraction (synchronous, thread-offloaded)
- pandas>=2.2.0: Excel/CSV data extraction (synchronous, thread-offloaded)
- openpyxl>=3.1.2: Excel file support for pandas

Last Grunted: 02/05/2026
"""
from __future__ import annotations

import asyncio
import logging
import os
from enum import Enum
from io import BytesIO
from pathlib import Path

from fastapi import UploadFile

logger = logging.getLogger(__name__)

SHARED_VOLUME_PATH = os.getenv('SHARED_VOLUME_PATH', '/shared-files')

# Minimum text length to consider PDF text extraction successful
PDF_TEXT_MIN_CHARS = 200

# Centralized config
from docproc.config import get_ingestion_config as _get_config


class ProcessingStrategy(Enum):
    """Document processing strategies."""
    TEXT_ONLY = "text"          # Direct text extraction
    STRUCTURED_DATA = "data"    # Pandas/tabular extraction  
    VISION_REQUIRED = "vision"  # Vision model extraction


def _detect_strategy(filename: str, visual_analysis: bool) -> ProcessingStrategy:
    """
    Determine optimal processing strategy for a file.
    
    Uses the centralized IngestionConfig to check pdf_always_use_vision.
    
    Args:
        filename: Original filename with extension
        visual_analysis: If True, force vision for visual documents
    
    Returns:
        ProcessingStrategy enum value
    """
    ext = Path(filename).suffix.lower()
    
    # Text files - always direct
    if ext in {'.txt', '.md', '.rst', '.json', '.xml', '.html'}:
        return ProcessingStrategy.TEXT_ONLY
    
    # Structured data - pandas handles this efficiently
    if ext in {'.xlsx', '.xls', '.csv', '.tsv'}:
        return ProcessingStrategy.STRUCTURED_DATA
    
    # Word documents - python-docx
    if ext in {'.docx'}:
        return ProcessingStrategy.TEXT_ONLY
    
    # PPTX - always vision (layout matters)
    if ext == '.pptx':
        return ProcessingStrategy.VISION_REQUIRED
    
    # PDF - route based on config and flags
    if ext == '.pdf':
        cfg = _get_config()
        if visual_analysis or cfg.pdf_always_use_vision:
            return ProcessingStrategy.VISION_REQUIRED
        return ProcessingStrategy.TEXT_ONLY
    
    # Unknown - try text
    return ProcessingStrategy.TEXT_ONLY


async def read_file_from_shared_volume(file_ref) -> bytes:
    """
    Read file from shared volume using FileRef.

    Handles various path formats from different services. File I/O is
    offloaded to a thread to avoid blocking the event loop.

    Args:
        file_ref: FileRef object with path attribute

    Returns:
        File contents as bytes

    Raises:
        FileNotFoundError: If file cannot be located

    Last Grunted: 02/05/2026
    """
    def _sync_read() -> bytes:
        original_path = file_ref.path

        # Normalize path - handle various prefix formats
        if original_path.startswith('/shared/files/'):
            rel_path = original_path.replace('/shared/files/', '', 1)
            file_path = os.path.join(SHARED_VOLUME_PATH, rel_path)
        elif original_path.startswith('/shared-files/'):
            file_path = original_path
        elif original_path.startswith('/'):
            file_path = original_path
        else:
            file_path = os.path.join(SHARED_VOLUME_PATH, original_path)

        # Try primary path first
        if os.path.exists(file_path):
            with open(file_path, 'rb') as f:
                return f.read()

        # Try alternatives
        alternatives = [
            os.path.join(SHARED_VOLUME_PATH, os.path.basename(original_path)),
            os.path.join(SHARED_VOLUME_PATH, original_path.lstrip('/')),
        ]

        for alt_path in alternatives:
            if os.path.exists(alt_path):
                logger.debug(f"Found file at alternative path: {alt_path}")
                with open(alt_path, 'rb') as f:
                    return f.read()

        raise FileNotFoundError(f"File not found: {original_path}")

    return await asyncio.to_thread(_sync_read)


async def extract_text_from_file(
    file: UploadFile,
    visual_analysis: bool = False,
    model: str = "gpt-4o-mini",
) -> tuple[str, str, str | None, str]:
    """
    Extract text from uploaded file using smart routing.
    
    Automatically selects the best extraction method based on file type:
    - Text/Markdown: Direct UTF-8 read (with error replacement)
    - XLSX/CSV: Pandas extraction with multi-sheet support
    - DOCX: python-docx extraction (paragraphs + tables)
    - PDF: pdfplumber text extraction first, vision fallback if < 200 chars
    - PPTX: Vision extraction (layout/positioning matters)
    
    Args:
        file (UploadFile): FastAPI UploadFile object with filename and read() method
        visual_analysis (bool): If True, include visual descriptions for visual docs.
            Forces vision extraction for PDFs regardless of text content.
        model (str): LLM model for vision extraction (default: "gpt-4o-mini")
    
    Returns:
        tuple[str, str, str | None, str]: (filename, text, error, visual_info)
            - filename: Original filename (stripped of whitespace)
            - text: Extracted text content (may be empty on failure)
            - error: Error message if failed, else None
            - visual_info: Empty string (visual descriptions now inline in text)
    
    Note:
        PDF text extraction threshold: 200 characters minimum to be considered
        successful. Below this, falls back to vision extraction.
    
    Last Grunted: 02/05/2026
    """
    filename = file.filename.strip()
    ext = Path(filename).suffix.lower()
    
    try:
        content = await file.read()
        strategy = _detect_strategy(filename, visual_analysis)
        
        # Route to appropriate extractor
        if strategy == ProcessingStrategy.TEXT_ONLY:
            if ext == '.docx':
                text = await _extract_docx(content)
            else:
                text = content.decode('utf-8', errors='ignore')
            return filename, text, None, ""
        
        if strategy == ProcessingStrategy.STRUCTURED_DATA:
            text = await _extract_structured_data(content, ext)
            return filename, text, None, ""
        
        if strategy == ProcessingStrategy.VISION_REQUIRED:
            cfg = _get_config()
            # For PDFs, try pdfplumber first UNLESS vision is forced
            if ext == '.pdf' and not visual_analysis and not cfg.pdf_always_use_vision:
                text = await _extract_pdf_text(content)
                if len(text.strip()) >= PDF_TEXT_MIN_CHARS:
                    logger.info(f"PDF {filename}: text extraction successful ({len(text)} chars)")
                    return filename, text, None, ""
                logger.info(f"PDF {filename}: insufficient text, using vision")
            elif ext == '.pdf' and cfg.pdf_always_use_vision:
                logger.info(f"PDF {filename}: using vision (pdf_always_use_vision=True)")
            
            # Use vision extraction
            from docproc.services.multimodal_extractor import extract_document_with_vision
            text, visual_info, error = await extract_document_with_vision(
                content, ext, filename, visual_analysis
            )
            return filename, text, error, visual_info
        
        return filename, "", f"Unsupported file type: {ext}", ""
        
    except Exception as exc:
        logger.exception(f"Extraction failed for {filename}: {exc}")
        return filename, "", str(exc), ""


async def download_and_process_url(
    url: str,
    visual_analysis: bool = True,
    model: str = "gpt-4o-mini",
) -> tuple[str, str, str | None, str]:
    """
    Download content from URL and extract text.
    
    Supports PDF, PPTX, DOCX, XLSX, and text files.
    
    Args:
        url: URL to download from
        visual_analysis: Include visual descriptions for visual docs
        model: LLM model for vision extraction
    
    Returns:
        Tuple of (filename, text, error, visual_info)
    
    Last Grunted: 02/05/2026
    """
    import aiohttp
    from urllib.parse import urlparse, unquote
    
    try:
        parsed = urlparse(url)
        path = unquote(parsed.path)
        filename = os.path.basename(path) if path else "downloaded_content"
        
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=aiohttp.ClientTimeout(total=120)) as response:
                if response.status != 200:
                    return filename, "", f"HTTP {response.status}", ""
                
                content = await response.read()
                content_type = response.headers.get('Content-Type', '').lower()
                
                # Infer extension from content type if needed
                ext = _infer_extension(filename, content_type)
                if not filename.endswith(ext):
                    filename = f"{Path(filename).stem}{ext}"
        
        # Create mock UploadFile-like object
        class MockFile:
            def __init__(self, name: str, data: bytes):
                self.filename = name
                self._data = data
            
            async def read(self):
                return self._data
        
        mock_file = MockFile(filename, content)
        return await extract_text_from_file(mock_file, visual_analysis, model)
        
    except asyncio.TimeoutError:
        return "unknown", "", "Download timeout (120s)", ""
    except Exception as exc:
        logger.exception(f"URL download failed: {url}")
        return "unknown", "", str(exc), ""


def _infer_extension(filename: str, content_type: str) -> str:
    """Infer file extension from content type."""
    ext = Path(filename).suffix.lower()
    if ext:
        return ext
    
    type_map = {
        'pdf': '.pdf',
        'powerpoint': '.pptx',
        'presentation': '.pptx',
        'spreadsheet': '.xlsx',
        'excel': '.xlsx',
        'word': '.docx',
        'text/plain': '.txt',
        'text/csv': '.csv',
    }
    
    for key, value in type_map.items():
        if key in content_type:
            return value
    
    return '.pdf'  # Default assumption


async def _extract_pdf_text(content: bytes) -> str:
    """
    Extract text from PDF using pdfplumber (no vision API needed).

    This is the fast path for text-based PDFs. Uses pdfplumber's extract_text()
    method which works best with machine-generated PDFs. Falls back to empty
    string if extraction fails, allowing vision fallback.

    pdfplumber is synchronous and can be CPU-intensive for large PDFs, so the
    extraction is offloaded to a thread via asyncio.to_thread().

    Args:
        content (bytes): Raw PDF file bytes

    Returns:
        str: Extracted text with page markers (e.g., "[Page 1]\\n...").
            Empty string if extraction fails or pdfplumber not installed.

    Limits:
        - Max 50 pages processed to avoid memory issues
        - Each page's text is prefixed with [Page N]

    Note:
        pdfplumber (v0.11.9 verified) builds on pdfminer.six and works best
        with machine-generated PDFs rather than scanned documents.

    Last Grunted: 02/05/2026
    """
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber not installed, skipping text extraction")
        return ""

    def _sync_extract() -> str:
        buffer = BytesIO(content)
        text_parts: list[str] = []
        with pdfplumber.open(buffer) as pdf:
            for i, page in enumerate(pdf.pages[:50], 1):  # Limit pages
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(f"[Page {i}]\n{page_text}")
        return "\n\n".join(text_parts)

    try:
        return await asyncio.to_thread(_sync_extract)
    except Exception as exc:
        logger.warning(f"PDF text extraction failed: {exc}")
        return ""


async def _extract_docx(content: bytes) -> str:
    """
    Extract text from DOCX using python-docx.

    Extracts all paragraphs and tables from the document. Tables are
    converted to pipe-separated format for readability. python-docx is
    synchronous, so extraction is offloaded to a thread.

    Args:
        content (bytes): Raw DOCX file bytes (Office Open XML format)

    Returns:
        str: Extracted text with paragraphs and tables joined by double newlines.

    Raises:
        RuntimeError: If python-docx is not installed
        Exception: Re-raises any extraction errors

    Note:
        python-docx (v1.1.0 verified) extracts only text content.
        Images, charts, and other embedded objects are not extracted.

    Last Grunted: 02/05/2026
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed")
        raise RuntimeError("python-docx required for DOCX files. pip install python-docx")

    def _sync_extract() -> str:
        buffer = BytesIO(content)
        doc = Document(buffer)
        text_parts: list[str] = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(" | ".join(row_text))
            if table_text:
                text_parts.append("\n".join(table_text))

        return "\n\n".join(text_parts)

    try:
        return await asyncio.to_thread(_sync_extract)
    except Exception as exc:
        logger.exception(f"DOCX extraction failed: {exc}")
        raise


async def _extract_structured_data(content: bytes, ext: str) -> str:
    """
    Extract text from structured data files (XLSX, CSV) using pandas.

    Converts spreadsheet data to readable text format suitable for
    summarization and embedding. Excel files with multiple sheets
    are processed with sheet name headers. pandas is synchronous, so
    extraction is offloaded to a thread.

    Args:
        content (bytes): Raw file bytes
        ext (str): File extension including dot (e.g., ".xlsx", ".csv")

    Returns:
        str: Text representation of the data. Large datasets (>500 rows)
            are truncated with a note.

    Raises:
        RuntimeError: If pandas/openpyxl not installed
        Exception: Re-raises any extraction errors

    Supported formats:
        - .xlsx, .xls: Excel files (requires openpyxl)
        - .csv: Comma-separated values
        - .tsv: Tab-separated values

    Note:
        pandas (v2.2.0 verified) with openpyxl backend for Excel support.

    Last Grunted: 02/05/2026
    """
    try:
        import pandas as pd
    except ImportError:
        logger.error("pandas not installed")
        raise RuntimeError("pandas required for spreadsheet files. pip install pandas openpyxl")

    def _sync_extract() -> str:
        buffer = BytesIO(content)
        text_parts: list[str] = []

        if ext in {'.xlsx', '.xls'}:
            # Excel file - may have multiple sheets
            xlsx = pd.ExcelFile(buffer)
            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)
                if not df.empty:
                    text_parts.append(f"=== Sheet: {sheet_name} ===")
                    text_parts.append(_dataframe_to_text(df))

        elif ext in {'.csv', '.tsv'}:
            # CSV/TSV - single sheet
            sep = '\t' if ext == '.tsv' else ','
            df = pd.read_csv(buffer, sep=sep)
            text_parts.append(_dataframe_to_text(df))

        return "\n\n".join(text_parts)

    try:
        return await asyncio.to_thread(_sync_extract)
    except Exception as exc:
        logger.exception(f"Structured data extraction failed: {exc}")
        raise


def _dataframe_to_text(df) -> str:
    """Convert pandas DataFrame to readable text."""
    original_len = len(df)

    # Limit rows for very large datasets
    if original_len > 500:
        df = df.head(500)
        truncated = True
    else:
        truncated = False

    # Convert to string representation
    text = df.to_string(index=False, max_colwidth=100)

    if truncated:
        text += f"\n\n[... showing 500 of {original_len} rows, data truncated ...]"

    return text
